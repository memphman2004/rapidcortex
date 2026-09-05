#!/usr/bin/env python3
"""Align IU Component C SOW + narrative with Rapid Cortex product facts."""
from __future__ import annotations

import shutil
from copy import copy
from pathlib import Path

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC_XLSX = Path("/Users/jeffcoleman/Downloads/Rapid_Cortex_IU_Component_C_SOW_Response_Completed.xlsx")
SRC_DOCX = Path("/Users/jeffcoleman/Downloads/RC_IU_RFP_TEC_1952_2027_Component_C-2.docx")
OUT_DIR = Path("/Users/jeffcoleman/Downloads")
OUT_XLSX = OUT_DIR / "Rapid_Cortex_IU_Component_C_SOW_Response_Aligned.xlsx"
OUT_DOCX = OUT_DIR / "RC_IU_RFP_TEC_1952_2027_Component_C-2_Aligned.docx"

# Spreadsheet dropdown values (do not invent others on SOC / platform sheets)
FULLY = "Fully Met"
PARTIAL = "Partially Met"
ROADMAP = "Roadmap"
NA = "Not Applicable"

EXCEL_UPDATES = {
    "SOC-001": (
        PARTIAL,
        "Rapid Cortex is an API-first, vendor-agnostic situational-awareness layer. REST/JSON APIs, webhooks, and normalized incident/event ingest are in production today. Device-layer protocols (ONVIF, OSDP, SIA, MQTT, Wiegand) remain with IU’s source systems (Milestone XProtect, CBORD/CS Gold, sensors) or are delivered as Year-1 connectors; Rapid Cortex does not terminate those protocols itself. IU keeps best-of-breed subsystems; Rapid Cortex unifies incidents, maps, audit, and operator workflows without a single-vendor lock-in.",
    ),
    "SOC-007": (
        PARTIAL,
        "Rapid Cortex natively associates cameras, QR/NFC locations, buildings, floors, and zones to geospatial records and uses that mapping for nearest-resource and incident context. Mapping emergency phones (Component B) or access points to cameras is IU configuration plus the relevant connector. The geographic association model is live; IU-specific phone-to-camera tables will be loaded during implementation.",
    ),
    "SOC-022": (
        PARTIAL,
        "The Rapid Cortex event model consolidates alerts from multiple sources into one queue, map, and dashboard with source attribution. Production ALPR connectors (Flock, Genetec AutoVu, or other IU-named platforms) are implementation work, not pre-certified adapters. Once connected, operators see a unified view; the ALPR vendor remains the system of record for plate recognition and hot-list administration.",
    ),
    "SOC-027": (
        PARTIAL,
        "Rapid Cortex provides a live, map-based campus and venue command picture (Mapbox) with incidents, QR/NFC/SMS intake, campus-scoped RBAC, and a nine-campus enterprise view for authorized staff. Responder GPS, EMNS delivery receipts, and live VMS wall status appear when those source systems provide them through Year-1 connectors. Rapid Cortex does not replace IU’s EMNS, officer AVL, or Milestone as systems of record.",
    ),
    "SOC-028": (
        PARTIAL,
        "The platform is built to aggregate heterogeneous public-safety events into one operational picture. Native today: QR/NFC/SMS reports, incident lifecycle, maps, audit, translation, AI-assisted classification, and Ring/KVS-linked live video. EMNS, blue-light phones, CBORD, Milestone, ALPR, and sensors join through contracted connectors while those products remain authoritative. This matches IU’s phased coexistence model.",
    ),
    "SOC-035": (
        PARTIAL,
        "Native incident history, filtering, time/location/type views, and command reporting are included across campuses in the IU tenant. CompStat-style heat maps, fusion-center ingest, and advanced cross-jurisdiction crime analytics are offered as the optional Advanced Crime Analytics module ($15,000/year), not as a current core module.",
    ),
    "SOC-040": (
        PARTIAL,
        "Rapid Cortex provides a unified incident/alarm queue with source type, priority, location, acknowledgment, ownership, escalation, and audit. QR/NFC panic, SMS SOS, and connected-system events are normalized when ingested. Microkey (or successor) fire/security alarm ingest is a Year-1 interface evaluation, not a current native connector.",
    ),
    "SOC-041": (
        PARTIAL,
        "Rapid Cortex associates SOPs, checklists, and response guidance with incident type and location, including campus protocol-style coaching for operators. A full IU-authored building EAP document library with automatic retrieval of PDF EAPs is configuration and content load during implementation, not a certified EAP-management product out of the box.",
    ),
    "SOC-042": (
        FULLY,
        "Rapid Cortex includes AI-assisted incident intelligence via AWS Bedrock (Anthropic Claude primary, with configured fallbacks): NLP for summarization and classification, live transcription, and 100-plus written translation language codes. Computer vision, gunshot, and VMS analytics are ingested from IU’s cameras/VMS when those systems emit events. AI is decision support with human override; it does not replace operator judgment or statutory reporting.",
    ),
    "SOC-043": (
        PARTIAL,
        "Rules and event-driven automation exist for incident creation, notifications, and escalation based on type, location, and severity. Weather (BAMWX), gunshot, CAD-type, and access-control auto-workflows depend on those Year-1 source connectors. Consequential actions (lockdown commands, CAD write) remain human-confirmable and fail-closed.",
    ),
    "INT-023": (
        PARTIAL,
        "Enterprise SSO is delivered through Amazon Cognito Hosted UI federation (SAML 2.0 and/or OpenID Connect) to IU’s IdP (Shibboleth, Azure AD, Duo). Custom in-app username/password routes are not the IU path. Federation, MFA policy, and JIT attributes will be configured and acceptance-tested during onboarding.",
    ),
    "UM-019": (
        PARTIAL,
        "Rapid Cortex uses a canonical RBAC model with predefined campus, venue, hospital, transit, and PSAP roles—including campus_admin, campus_supervisor, campus_security, campus_counselor, and campus_faculty—scoped by agencyId and campus. IU cannot invent arbitrary custom permission sets in a product role-builder; additional roles are contracted configuration, not an in-product policy editor.",
    ),
    "INT-026": (
        PARTIAL,
        "Native operational mapping is Mapbox. Campus/venue zones, incident pins, and overlays are in production. Esri ArcGIS or other GIS layers can be ingested as web services/overlays during implementation. Rapid Cortex is not an ArcGIS replacement.",
    ),
    "ITS-012": (
        PARTIAL,
        "Rapid Cortex is cloud SaaS. Priority 1 production incidents are acknowledged within 30 minutes via on-call escalation, with remote diagnosis as the primary remediation path. Rapid Cortex does not staff a 24/7 security operations center on IU’s behalf and does not dispatch field technicians or spare parts for IU hardware. The four-hour on-site hardware response in this requirement does not apply to Component C; device repair remains with EMNS, blue-light, VMS, and related vendors. Awareness-only sites receive the same remote software support.",
    ),
    "INT-027": (
        PARTIAL,
        "In-platform war rooms (collaborative incident command) are a Rapid Cortex product surface. Deep Microsoft Teams meeting/channel integration is an optional Year-1 module ($8,000/year) via Microsoft Graph APIs, subject to IU Teams telephony permissions. Slack can use equivalent webhook/API patterns.",
    ),
    "SOC-002": (
        FULLY,
        "Subsystem integrations are decoupled from the Rapid Cortex incident and operational data model. A camera, ALPR, access-control, or sensor source can be added or replaced by changing its connector while preserving incidents, workflows, dashboards, audit history, and campus configuration. Live today this is proven with Ring, AWS KVS, and RTSP/ONVIF-capable camera paths; Milestone/CBORD/ALPR follow the same connector pattern in Year 1.",
    ),
    "SOC-026": (
        FULLY,
        "Rapid Cortex provides native incident creation, assignment, status changes, notes, media, escalation, and closure with a unified chronological continuity log. Intake channels include QR, NFC, dedicated SMS, console entry, and webhooks. In-platform war rooms support major-incident coordination. Events from integrated systems attach to the same incident so operators keep a single operational record.",
    ),
    "TI-004": (
        FULLY,
        "Rapid Cortex uses logical tenant isolation: every DynamoDB access is scoped by agencyId, and JWT role/campus claims enforce RBAC. IU campuses/buildings/zones can be segmented so local operators see only authorized scope while enterprise Public Safety leadership can receive a consolidated nine-campus view. Tenants share AWS services (Lambda, tables, S3) with partition-level isolation; other customers cannot read IU records. URL slugs are not the security boundary.",
    ),
}

# Additional Excel explanation-only updates (keep rating)
EXCEL_EXPLAIN_ONLY = {
    "SOC-012": "Rapid Cortex is camera-vendor agnostic and consumes camera/VMS context rather than replacing the VMS. Hanwha Wisenet discovery, PTZ, and H.265 remain with Milestone XProtect / ONVIF-capable infrastructure. Direct RTSP/KVS paths exist for venue/campus live video; IU’s preferred architecture is Hanwha via XProtect, which is a Year-1 connector workstream.",
    "SOC-015": "Privacy masking of live/recorded imagery remains controlled by Milestone XProtect / Hanwha so the authoritative privacy configuration is preserved. Rapid Cortex enforces role-based access, campus/zone scoping, and audit around incident-linked video. It does not apply pixel masks inside a Rapid Cortex camera UI.",
    "SOC-032": "ARMS by End2End is a priority Year-1 integration. Rapid Cortex can ingest CAD incident type/location and exchange status through available ARMS interfaces. CAD write-back, if IU elects it, is fail-closed and requires human review; Rapid Cortex does not auto-create CAD records. Exact bi-directional functions depend on End2End interface availability and licensing.",
    "INT-025": "ARMS by End2End is a priority integration. Rapid Cortex can ingest CAD incident type/location and exchange incident/status data through available ARMS interfaces. CAD write-back is fail-closed and human-approved. Other CAD platforms can be integrated through documented APIs; Rapid Cortex does not replace ARMS as the CAD/RMS of record.",
    "UM-016": "SCIM 2.0 lifecycle provisioning is planned for the IU enterprise identity integration (Roadmap). Until enabled, provisioning is supported through Cognito SSO claims (JIT), APIs, and controlled bulk/scheduled synchronization.",
    "ITS-028": "Rapid Cortex will provide 24/7/365 on-call escalation for Priority 1 production software incidents with phone/on-call access, plus email/portal support for lower priorities. This is software on-call, not a staffed Rapid Cortex SOC operating IU’s cameras and doors. Final support channels and escalation contacts will be documented in the SLA.",
    "STD-011": "Rapid Cortex will document applicable security-framework alignment and AWS inherited controls in the HECVAT/security package. AWS services carry their own SOC 2 / ISO / FedRAMP attestations; those are not Rapid Cortex product certifications. Rapid Cortex does not claim ISO 27001, SAFETY Act, or SOC 2 Type II for the Rapid Cortex application unless separately evidenced at submission.",
}


def set_run_text(paragraph: Paragraph, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def replace_exact(paragraph: Paragraph, mapping: dict[str, str]) -> bool:
    text = paragraph.text
    if text in mapping:
        set_run_text(paragraph, mapping[text])
        return True
    return False


def replace_cell_text(cell, old: str, new: str) -> bool:
    joined = cell.text
    if old not in joined and joined != old:
        return False
    # Prefer exact paragraph match inside the cell
    for p in cell.paragraphs:
        if p.text == old or old in (p.text or ""):
            if p.text == old:
                set_run_text(p, new)
            else:
                set_run_text(p, (p.text or "").replace(old, new))
            return True
    if cell.paragraphs:
        set_run_text(cell.paragraphs[0], new)
        for p in cell.paragraphs[1:]:
            set_run_text(p, "")
        return True
    return False


def walk_tables(doc: Document):
    for table in doc.tables:
        yield table
        # nested tables
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    yield nested


def align_excel() -> None:
    shutil.copy2(SRC_XLSX, OUT_XLSX)
    wb = openpyxl.load_workbook(OUT_XLSX)
    by_req: dict[str, tuple] = {}
    for sheet in wb.worksheets:
        for row in range(1, sheet.max_row + 1):
            req = sheet.cell(row, 2).value
            if isinstance(req, str) and req.strip():
                by_req.setdefault(req.strip(), [])
                by_req[req.strip()].append((sheet, row))

    missing = []
    for req, payload in EXCEL_UPDATES.items():
        status, explanation = payload
        hits = by_req.get(req)
        if not hits:
            missing.append(req)
            continue
        for sheet, row in hits:
            sheet.cell(row, 6).value = status
            sheet.cell(row, 7).value = explanation
    for req, explanation in EXCEL_EXPLAIN_ONLY.items():
        hits = by_req.get(req)
        if not hits:
            missing.append(req)
            continue
        for sheet, row in hits:
            sheet.cell(row, 7).value = explanation
    if missing:
        raise SystemExit(f"Missing Excel req ids: {missing}")
    wb.save(OUT_XLSX)


PARA = {
    "Rapid Cortex fully understands Indiana University's phased implementation approach and has designed its proposal to deliver standalone value for Component C while maintaining clear integration paths to the full platform vision as described in Section F of the RFP.": (
        "Rapid Cortex is proposing the campus-native operational layer IU can run on day one while coexistence with incumbent systems continues. "
        "Native today: QR/NFC/SMS safety reporting (no public app download), campus and venue consoles, incident lifecycle with a unified timeline, "
        "Mapbox operational maps, campus-scoped RBAC (including counselor and faculty roles), AI-assisted classification via AWS Bedrock, "
        "live transcription, 100-plus written translation language codes, audit, Clery-adjacent classification/retention support, and Ring/KVS live video. "
        "Year-1 Founding Partner connectors: Milestone XProtect, CBORD/CS Gold, ARMS by End2End, and BAMWX. "
        "Rapid Cortex does not replace 911, CAD, Milestone, CBORD, Clery statutory filing, or IU’s chosen EMNS and blue-light vendors. "
        "Ratings in the attached SOW workbook are controlling; this narrative expands those ratings and is aligned to them."
    ),
    "The Rapid Cortex platform delivers real-time incident intelligence, multi-channel safety reporting via QR codes, NFC tags, and dedicated SMS lines, live transcription, bidirectional multilingual translation across 40-plus languages, AI-powered incident classification, supervisor command visibility, and integrated camera access — from a single unified cloud-native platform. The system is designed to operate alongside existing infrastructure without requiring replacement of any incumbent technology.": (
        "The Rapid Cortex platform delivers real-time incident intelligence; multi-channel safety reporting via QR codes, NFC tags, and dedicated SMS lines; "
        "live transcription; bidirectional written translation across 100-plus language codes (speech-to-text uses a configured subset); "
        "AI-assisted incident classification via AWS Bedrock; supervisor command visibility; and live camera access through VMS/KVS paths — "
        "from a single unified cloud-native platform. The system is designed to operate alongside existing infrastructure without requiring replacement of CAD, VMS, access control, EMNS, or blue-light systems."
    ),
    "The platform is architected on Amazon Web Services (AWS) and built with CJIS Security Policy alignment from the ground up. Core technology components include:": (
        "The platform is architected on Amazon Web Services (AWS) with CJIS-aligned technical controls (not a CJIS certification). Core technology components include:"
    ),
    "Microsoft Teams integration for incident command and war room coordination": (
        "In-platform war rooms for major-incident command; optional Microsoft Teams meeting/channel integration as a Year-1 module"
    ),
    "The Rapid Cortex platform is designed with Clery Act compliance requirements as a core architectural consideration. Relevant capabilities include:": (
        "Rapid Cortex supports Clery-adjacent operational recordkeeping. Clery Act statutory filings, Annual Security Reports, and Emergency Notification / Timely Warning issuance remain IU’s responsibility and, for mass notification, the Component A EMNS vendor. Relevant Rapid Cortex capabilities include:"
    ),
    "Incident classification workflows support Emergency Notification and Timely Warning distinctions with separate templates, approval workflows, and audit trails.": (
        "Incident classification can record Clery-relevant categories and geography to support IU’s Clery workflow. Rapid Cortex does not send Emergency Notifications or Timely Warnings; those remain IU policy plus the selected EMNS."
    ),
    "System testing documentation including date, time, scope, and delivery results suitable for Clery Act test documentation requirements.": (
        "Timestamped exercise and incident/audit records suitable for after-action documentation. Audible/channel delivery tests remain with the EMNS and blue-light systems; Rapid Cortex can log integrated test events when those systems notify it."
    ),
    "Rapid Cortex is designed with student data privacy as a first-order requirement. The platform does not store student education records and does not expose personally identifiable student information in operational interfaces without explicit role authorization. Specifically:": (
        "Rapid Cortex is designed with student data privacy as a first-order requirement. Safety reports submitted via QR, NFC, or SMS may contain personally identifiable information; those records are operational public-safety records, not a student-information system. FERPA applicability is an IU legal/policy determination under a Data Processing Agreement. Rapid Cortex will not use IU data for marketing. Specifically:"
    ),
    "Incident reports are routed by phone number or QR/NFC zone — not by student identity.": (
        "Incident reports are routed by dedicated SMS number, QR/NFC location, and campus scope — not by a SIS student record. Reporters may remain anonymous or identify themselves."
    ),
    "Privacy masking zones can be configured for camera streams near healthcare, counseling, or sensitive facilities.": (
        "Camera privacy masking for healthcare, counseling, or other sensitive facilities remains in Milestone XProtect / Hanwha. Rapid Cortex restricts who can open incident-linked video by role, campus, and audit."
    ),
    "All data access is logged with user ID, timestamp, and action for FERPA compliance audit purposes.": (
        "All data access is logged with user ID, timestamp, and action to support IU FERPA/privacy audits."
    ),
    "The Rapid Cortex anonymous tip reporting system supports confidential incident submission without requiring reporter identification. This enables students to report Title IX concerns through QR codes, NFC tags, or dedicated SMS lines without risk of identity exposure. The back-end workflow supports tip intake, assignment to appropriate staff, and anonymous follow-up with the reporter while preserving anonymity throughout the chain.": (
        "Anonymous QR/NFC/SMS intake supports confidential reporting without requiring reporter identification, including concerns that IU may later handle under Title IX. Reports can be assigned to campus counselor or security roles. Rapid Cortex is not a Title IX case-management system, investigation file, or hearing platform; IU’s Title IX office remains system of record for those processes."
    ),
    "Rapid Cortex delivers SMS through AWS End User Messaging as its primary carrier infrastructure — a direct AWS service with established relationships with all major US wireless carriers including AT&T, T-Mobile, Verizon, and regional carriers. This is not an aggregator model. AWS End User Messaging delivers messages through direct carrier connections, providing enterprise-grade reliability, throughput, and delivery transparency.": (
        "Rapid Cortex delivers two-way SMS through AWS End User Messaging as its primary path, with Twilio as a secondary provisioning path for campus/venue dedicated lines. AWS End User Messaging is an AWS-managed messaging service with US carrier reach. Rapid Cortex uses it for conversational incident reporting (Component C), not as IU’s mass-notification engine."
    ),
    "Carrier-level delivery receipts — per-message delivery status tracking to individual handsets": (
        "Delivery status as provided by AWS End User Messaging / Twilio (message accepted/failed at the provider). Handset-level carrier receipts are not independently guaranteed by Rapid Cortex."
    ),
    "AI intelligence processing (RapidIQ): $40,000/year": (
        "AI incident intelligence (AWS Bedrock): $40,000/year"
    ),
    "Encryption at rest: All DynamoDB tables and S3 buckets encrypted with AWS KMS customer-managed keys.": (
        "Encryption at rest: DynamoDB and S3 use AWS-managed encryption at rest (SSE). Customer-managed KMS keys can be applied when IU’s security plan requires them."
    ),
    "Penetration testing: Annual third-party penetration testing of the platform.": (
        "Penetration testing: Rapid Cortex will contract annual third-party penetration testing of the platform under the IU agreement and share results under NDA."
    ),
    "Rapid Cortex acknowledges and commits to compliance with all data types specified in Question 20(b), including FERPA, HIPAA (with BAA if PHI is processed), GLBA, Indiana SSN Law, Indiana Breach Notification Laws, and PCI-DSS. PCI-DSS is not applicable to the core RC Campus platform as the platform does not process payment card data. If any future module requires payment processing, a third-party PCI-compliant payment processor will be used.": (
        "Rapid Cortex will operate under an IU Data Processing Agreement and align technical and organizational controls to IU’s requirements for FERPA, applicable Indiana privacy and breach-notification laws, and GLBA where IU determines those laws apply to the data in this system. HIPAA: a BAA in IU-approved form will be executed before any PHI is processed; Rapid Cortex does not claim HIPAA certification. PCI-DSS is not applicable to Component C (no cardholder data). Rapid Cortex does not claim FERPA, HIPAA, CJIS, or SOC 2 certification in this proposal."
    ),
    "ONVIF Profile S/G compliance for camera discovery and streaming": (
        "ONVIF/RTSP camera paths where IU’s VMS or cameras expose them; Hanwha discovery/PTZ remain with Milestone XProtect unless IU elects a direct stream"
    ),
    "MQTT support for IoT sensor data ingestion": (
        "IoT/sensor ingest via REST, webhooks, or MQTT gateways provided by the source system or a Year-1 connector — Rapid Cortex does not require IU to re-wire sensors"
    ),
    "SOC-002: Individual subsystem vendors can be added or replaced without platform migration. Each integration is a separately deployable connector module. Replacing a camera vendor requires reconfiguring the camera connector — not rebuilding the platform. This has been validated with RC's own multi-vendor camera architecture (Ring, AWS KVS, ONVIF-compatible cameras).": (
        "SOC-002: Individual subsystem vendors can be added or replaced without platform migration. Each integration is a separately deployable connector. Replacing a camera vendor requires reconfiguring the camera connector — not rebuilding incidents, maps, or RBAC. This pattern is live with Ring, AWS KVS, and RTSP/ONVIF-capable cameras; Milestone XProtect follows the same model in Year 1."
    ),
    "SOC-012 — Hanwha Wisenet: RC's ONVIF-compliant camera integration supports camera discovery, streaming (RTSP), and PTZ control for any ONVIF Profile S, G, or T camera — including the Hanwha Wisenet product line. Hanwha cameras can be integrated directly into the RC camera management layer without the Milestone XProtect intermediary for use cases where direct ONVIF streaming is preferred. Hanwha's Wisenet WAVE and Wisenet SSM platforms also support open API access that RC can leverage for deeper integration.": (
        "SOC-012 — Hanwha Wisenet: Rapid Cortex is not a camera VMS. IU’s preferred path is Hanwha cameras managed in Milestone XProtect, with live/recorded context presented in the Rapid Cortex SOC console through the Year-1 XProtect connector. Where IU provides an RTSP/ONVIF URL, Rapid Cortex can stream via AWS KVS/WebRTC for incident-linked live video. PTZ, Profile S/G/T discovery, and H.265 multi-stream remain XProtect/Hanwha functions."
    ),
    "SOC-013: RC supports H.264 and H.265 video streams via WebRTC and HLS delivery paths. Multi-stream configurations (high-res recording + low-res live preview) are supported. Edge recording is not a native RC function — RC streams live video; recording is handled by the VMS (Milestone XProtect).": (
        "SOC-013: Rapid Cortex can present H.264/H.265 streams made available by the VMS or a direct RTSP/KVS path via WebRTC. Multi-stream and edge recording remain Milestone/Hanwha functions. Rapid Cortex does not transcode IU’s camera fleet or replace XProtect recording."
    ),
    "SOC-015: Privacy masking zones are configurable per camera stream in the RC camera management interface. Masking is applied at the display layer within the RC console — cameras near healthcare, counseling, or FERPA-sensitive facilities can be restricted by role. HIPAA-sensitive camera streams can be limited to users with explicit permission grants.": (
        "SOC-015: Privacy masking of camera imagery is configured in Milestone XProtect / Hanwha so IU’s authoritative privacy zones are preserved. Rapid Cortex restricts who may open incident-linked video by role and campus, and audits that access. Rapid Cortex does not apply pixel masks in a Rapid Cortex camera-management UI."
    ),
    "SOC-017: RC supports vendor-agnostic access control integration via REST API (preferred), OSDP, and legacy Wiegand where required for backward compatibility. Supported or planned platforms include CBORD/CS Gold (Year 1), Lenel S2, Genetec Synergis, Honeywell Pro-Watch, and Software House C-CURE. All access events are normalized into RC's standard event schema before display.": (
        "SOC-017: Rapid Cortex integrates access control through REST/webhooks/middleware. OSDP and Wiegand remain device/controller-layer protocols terminated by CBORD (or another ACS), not by Rapid Cortex. CBORD/CS Gold is the Year-1 IU connector. Additional ACS vendors can follow the same event schema; they are not represented as currently certified connectors."
    ),
    "War room creation for major incidents with collaborative task management and Teams meeting integration": (
        "In-platform war rooms for major incidents with collaborative task management; Microsoft Teams meeting integration is an optional Year-1 module"
    ),
    "Push RC incident data to ARMS for CAD record creation": (
        "Share Rapid Cortex incident context with ARMS under IU-approved, human-reviewed workflows — CAD write-back is fail-closed and is not automatic"
    ),
    "SOC-039: NFC beacon-based patrol management is a core RC capability used for campus and venue security patrol tracking. RC generates unique NFC tags per patrol checkpoint — officers tap the tag on rounds, and check-ins are logged to the supervisor dashboard with timestamp and location. RC does not currently have a direct integration with Allied Universal's Heliaus platform, but Heliaus event data can be ingested via webhook to display patrol status in the RC console.": (
        "SOC-039: Rapid Cortex natively uses QR codes and NFC tags for public safety reporting and location context (buildings, floors, zones, venue sections). Those same tags can be used as officer checkpoint taps if IU configures that workflow. Rapid Cortex is not a replacement for Allied Universal Heliaus patrol management. Heliaus events can be ingested via vendor-supported interfaces/webhooks during implementation."
    ),
    "SOC-042 — AI/ML Features (RapidIQ): Rapid Cortex's RapidIQ AI engine is the most significant capability differentiator in this proposal. RapidIQ is a proprietary multi-model AI inference pipeline built on Anthropic Claude (via AWS Bedrock) as the primary model, with OpenAI GPT as secondary fallback and direct Anthropic API as tertiary. RapidIQ provides:": (
        "SOC-042 — AI/ML Features: Rapid Cortex incident intelligence is a multi-model pipeline on AWS Bedrock (Anthropic Claude primary, with configured OpenAI and other fallbacks). This is operator decision support — not an autonomous dispatch engine. Capabilities include:"
    ),
    "No competitor in the SOC/situational awareness category offers a comparable AI intelligence layer. RapidIQ is the only public safety AI engine that extracts structured incident intelligence from multi-source inputs in real time without dispatcher entry — converting raw reports into actionable operational intelligence automatically.": (
        "This AI layer is a practical differentiator for campus SOC operations: unstructured QR/SMS/NFC reports become structured, translated, classified incident records for human review in seconds. Operators can override any classification. Rapid Cortex does not claim to be the only public-safety AI product in the market."
    ),
    "SOC-044 — AI Governance: RapidIQ incorporates the following human-in-the-loop controls:": (
        "SOC-044 — AI Governance: Rapid Cortex incident intelligence incorporates the following human-in-the-loop controls:"
    ),
    "INT-014 — PSIM/VMS: RC serves as the PSIM layer for the RC Campus deployment. VMS integration via Milestone XProtect is committed for Year 1 as described in SOC-003 through SOC-010.": (
        "INT-014 — PSIM/VMS: Rapid Cortex is the unifying situational-awareness and incident layer for Component C. It is not a traditional PSIM appliance and does not replace Milestone XProtect. XProtect integration is a Year-1 Founding Partner commitment as described in SOC-003 through SOC-010."
    ),
    "IU's data is logically isolated from all other RC customers at the DynamoDB record level. No shared tables, shared compute, or shared storage exist between customer tenants.": (
        "IU’s data is logically isolated from other customers at the DynamoDB record level using agencyId partition keys and JWT-enforced RBAC. Tenants share AWS services (Lambda, DynamoDB tables, S3) with strict tenant isolation; other customers cannot read IU records. URL path segments are not the security boundary."
    ),
    "UM-015 — Automated Provisioning: RC supports automated user provisioning from authoritative sources via SCIM 2.0 and SAML 2.0. For Indiana University, this means users authenticated through IU's Shibboleth or Azure AD identity provider can be automatically provisioned in RC on first sign-in (JIT provisioning) with role and campus assignment driven by IU directory attributes.": (
        "UM-015 — Automated Provisioning: Rapid Cortex can provision users from IU’s IdP through Cognito federation and just-in-time account creation on first SSO sign-in, with role and campus assignment driven by agreed directory attributes. SIS/HRMS-specific mappings are implementation work. SCIM 2.0 lifecycle provisioning is Roadmap (see UM-016)."
    ),
    "UM-016 — SCIM 2.0: RC supports SCIM 2.0 for automated user lifecycle management including create, update, and deactivate operations. When a student or staff member is deprovisioned in IU's identity provider, their RC account is automatically deactivated.": (
        "UM-016 — SCIM 2.0: SCIM 2.0 lifecycle provisioning (create/update/deactivate) is planned for IU’s enterprise identity integration. Until SCIM is enabled, deprovisioning is handled through SSO session/claims, administrative disable, and scheduled directory sync. This matches the SOW rating of Roadmap."
    ),
    "UM-017 — Sync Frequency: User data synchronization is event-driven via SCIM webhooks for near-real-time provisioning changes. Batch sync is also available on a configurable schedule (daily minimum).": (
        "UM-017 — Sync Frequency: Synchronization can be event-driven where the source supports it, near-real-time through APIs, or scheduled/batch. Frequency is configurable based on source limits. SCIM webhooks are part of the SCIM roadmap, not a current dependency for go-live."
    ),
    "UM-019 — RBAC: RC provides nine pre-defined roles with granular, configurable permissions. For Indiana University, the relevant roles include:": (
        "UM-019 — RBAC: Rapid Cortex uses predefined canonical roles (not an arbitrary permission-set builder). Campus-relevant roles include:"
    ),
    "Custom roles with granular permission sets can be configured by RC for IU-specific requirements.": (
        "Additional roles or permission changes beyond the canonical set are contracted configuration performed by Rapid Cortex with IU, not a self-serve custom-role editor in the product."
    ),
    "INT-023 — SSO/SAML: RC is fully compatible with SAML 2.0 and supports OpenID Connect (OIDC) for federated identity. Integration with IU's Shibboleth identity provider and Azure AD / Duo SSO is supported and will be configured during onboarding.": (
        "INT-023 — SSO/SAML: IU users will authenticate through Amazon Cognito Hosted UI federated to IU’s IdP using SAML 2.0 and/or OpenID Connect (Shibboleth, Azure AD, Duo). MFA is enforced by IU’s IdP / Cognito policy. This is implementation work, rated Partially Met in the SOW workbook, and will be acceptance-tested during onboarding."
    ),
    "INT-024 — LDAP/AD: RC supports LDAP and Microsoft Active Directory integration for user directory synchronization via SCIM 2.0 bridge or direct LDAP connector.": (
        "INT-024 — LDAP/AD: Directory integration is through Microsoft/Azure APIs, SSO claims, or a securely brokered connector. Direct legacy LDAP and a SCIM bridge are selected only where required. SCIM itself remains Roadmap."
    ),
    "INT-026 — GIS: RC uses Mapbox for geographic incident visualization. ArcGIS integration is available via the ArcGIS REST API for GIS-based targeting and campus map overlays. Google Maps API is also supported.": (
        "INT-026 — GIS: Native operational mapping is Mapbox (campus/venue zones and incident visualization). Esri ArcGIS or other GIS layers can be ingested as overlays/web services during implementation. Rapid Cortex is not an ArcGIS or Google Maps replacement."
    ),
    "INT-027 — Microsoft Teams: Rapid Cortex has an active Microsoft Teams integration built on the Microsoft Graph API. The integration supports:": (
        "INT-027 — Microsoft Teams: In-platform Rapid Cortex war rooms are available for incident command. Deep Microsoft Teams meeting and channel integration is an optional Year-1 module ($8,000/year) using Microsoft Graph APIs, subject to IU licensing and admin consent. Planned capabilities:"
    ),
    "War room creation in RC automatically generates a Teams meeting — participants receive a join link directly in the RC war room interface": (
        "Optional: war room creation can generate a Teams meeting join link in the Rapid Cortex war room interface"
    ),
    "This is a live, production-ready integration.": (
        "Teams meeting/channel integration is optional scoped work, not a current core-platform dependency. In-platform war rooms do not require Teams."
    ),
    "ITS-036 — System Testing: RC supports monthly silent tests, semester audible tests, and annual comprehensive exercises. All test results are automatically logged with date, time, scope, channels tested, and results — formatted for Clery Act documentation requirements.": (
        "ITS-036 — System Testing: Rapid Cortex supports scheduled tabletop/functional exercises and captures timestamped incident/event/audit records for after-action review. Monthly silent tests and semester audible tests of mass-notification or blue-light hardware remain with those Component A/B systems. Rapid Cortex can document integrated test events when those systems notify it."
    ),
    "Describe the incident in their own language — RC automatically translates reports from 40-plus languages into English for the security console.": (
        "Describe the incident in their own language — Rapid Cortex translates written reports across 100-plus language codes into English for the security console (speech/STT uses a configured subset)."
    ),
    "The RC campus security console receives the report in real time, classified by RapidIQ AI (incident type, priority, recommended response), and assigned to the appropriate officer. The full exchange — report, follow-up messages, officer notes, and resolution — is logged to a single incident timeline auditable for Clery Act purposes.": (
        "The campus security console receives the report in real time, classified by Rapid Cortex incident intelligence (type, priority, recommended response) for human review, and routed to the appropriate campus_security, campus_supervisor, or campus_counselor role. The full exchange — report, follow-up messages, officer notes, and resolution — is logged to a single incident timeline that can support IU’s Clery-adjacent recordkeeping. Statutory Clery filings remain IU’s."
    ),
    "We recognize that Rapid Cortex is an emerging company without the institutional reference list that a 20-year incumbent vendor can present. We have been direct about this throughout this proposal. What we offer instead is a platform that is technically superior on every AI intelligence dimension, commercially structured to align our success with IU's success, and backed by a team that is deeply committed to making Indiana University the model for AI-native campus safety in the Big Ten.": (
        "We recognize that Rapid Cortex is an emerging company without the institutional reference list that a 20-year incumbent vendor can present. We have been direct about this throughout this proposal. What we offer instead is a campus-native SOC overlay that is already built for QR/NFC/SMS intake, multi-campus RBAC, AI-assisted operations, and coexistence with Milestone, CBORD, ARMS, and IU’s chosen EMNS — commercially structured as a Founding Partner so IU’s success is our reference, not a slide-deck promise."
    ),
    "SOC-028 — Unified Dashboard: The RC SOC dashboard aggregates events from all connected subsystems into a single operational picture. For IU, this will include: QR/NFC/SMS incident reports, access control events (CBORD), VMS camera triggers (Milestone), ALPR alerts, EMNS activation status, and environmental sensor alerts. Each event type has a distinct visual treatment and is filterable by source, campus, severity, and time window.": (
        "SOC-028 — Unified Dashboard: The Rapid Cortex SOC dashboard is built to aggregate connected subsystems into one operational picture. Native on day one: QR/NFC/SMS reports, incidents, maps, and Ring/KVS-linked video. CBORD, Milestone, ALPR, EMNS activation status, and sensors appear as those Year-1 connectors come online. Each event type is filterable by source, campus, severity, and time window. Source systems remain authoritative."
    ),
    "Computer vision: Live camera feed analysis for object detection and activity classification via connected VMS analytics platforms.": (
        "Computer vision: object detection and activity classification from IU’s VMS/camera analytics (ingested events), not a Rapid Cortex-native CV engine."
    ),
    "SSO configuration (SAML 2.0 connection to IU's Shibboleth/Azure AD identity provider)": (
        "SSO configuration via Amazon Cognito Hosted UI federated to IU’s Shibboleth/Azure AD/Duo identity provider (SAML 2.0 and/or OIDC)"
    ),
    "Records retention configurable to meet Clery Act's seven-year minimum requirement for all alert and incident records.": (
        "Records retention configurable to IU policy, including a seven-year setting to support Clery-adjacent incident records. Rapid Cortex does not file the Annual Security Report."
    ),
}


TABLE_CELL_REPLACEMENTS = [
    ("Hospital Security", "Hospital Routing"),
    ("Incident intake, staff coordination, escalation workflows", "Hospital capacity portal, routing, and staff coordination — not a hospital PSIM"),
    ("Multi-agency incident command, war rooms, stakeholder reporting", "Cross-agency incident sharing, in-platform war rooms, stakeholder reporting (feature-configured)"),
    ("AWS Partner Network", "AWS platform"),
    ("Audit engagement initiated; anticipated completion Q2 2027", "Controls designed for SOC 2 Type II; formal audit targeted, not completed. No SOC 2 report is claimed in this proposal."),
    ("Twilio Partner", "Twilio integration"),
    ("supervisor / commsupervisor", "supervisor"),
    ("AWS infrastructure: SOC 2 Type II, ISO 27001, FedRAMP (Moderate). RC platform: SOC 2 in progress (Q2 2027).", "AWS inherited controls (SOC 2 / ISO / FedRAMP on AWS services). Rapid Cortex application: SOC 2 Type II not completed; HECVAT and alignment documentation will be provided."),
    ("In Progress", "Planned"),  # careful - might hit other cells; we filter
]

CERT_TABLE_EXACT = {
    "Active": None,  # handled per-row below
}


def align_docx() -> None:
    shutil.copy2(SRC_DOCX, OUT_DOCX)
    doc = Document(str(OUT_DOCX))

    # Paragraphs
    for p in doc.paragraphs:
        replace_exact(p, PARA)

    # Tables: verticals, certifications, UM roles, STD matrix
    for table in walk_tables(doc):
        for row in table.rows:
            texts = [c.text.strip() for c in row.cells]
            # Certifications table rows
            if texts and texts[0] == "AWS Partner Network":
                replace_cell_text(row.cells[0], "AWS Partner Network", "AWS hosting")
                if len(row.cells) > 1:
                    replace_cell_text(row.cells[1], row.cells[1].text, "Active")
                if len(row.cells) > 2:
                    replace_cell_text(
                        row.cells[2],
                        row.cells[2].text,
                        "Production platform deployed on AWS (US-East-1). Not claimed as an AWS Partner Network membership benefit.",
                    )
            if texts and texts[0] == "SOC 2 Type II":
                if len(row.cells) > 1:
                    replace_cell_text(row.cells[1], row.cells[1].text, "Planned")
                if len(row.cells) > 2:
                    replace_cell_text(
                        row.cells[2],
                        row.cells[2].text,
                        "Application SOC 2 Type II report is not completed. Controls are designed for a future Type II audit (target window Q2 2027). Alignment documentation available; no certification is claimed.",
                    )
            if texts and texts[0] == "Twilio Partner":
                replace_cell_text(row.cells[0], "Twilio Partner", "Twilio integration")
                if len(row.cells) > 2:
                    replace_cell_text(
                        row.cells[2],
                        row.cells[2].text,
                        "Secondary SMS path for campus/venue dedicated 10DLC reporting lines. Not claimed as a Twilio Partner Program membership.",
                    )
            if texts and texts[0] == "Hospital Security":
                replace_cell_text(row.cells[0], "Hospital Security", "Hospital Routing")
                if len(row.cells) > 1:
                    replace_cell_text(
                        row.cells[1],
                        row.cells[1].text,
                        "Hospital capacity / routing portal and staff coordination — not a hospital PSIM or clinical EHR",
                    )
            if texts and texts[0] == "Emergency Management":
                if len(row.cells) > 1:
                    replace_cell_text(
                        row.cells[1],
                        row.cells[1].text,
                        "Cross-agency incident sharing, in-platform war rooms, stakeholder reporting",
                    )
            if texts and texts[0] == "Campus Safety":
                if len(row.cells) > 1:
                    replace_cell_text(
                        row.cells[1],
                        row.cells[1].text,
                        "QR/NFC/SMS reporting, campus console, counselor vs security routing, Clery-adjacent records, AI classification",
                    )
            if texts and texts[0] == "Venue Security":
                if len(row.cells) > 1:
                    replace_cell_text(
                        row.cells[1],
                        row.cells[1].text,
                        "Game-day zone reporting, section routing, capacity/cameras, Memorial Stadium-style venue console",
                    )
            if texts and texts[0] == "911 / PSAP":
                if len(row.cells) > 1:
                    replace_cell_text(
                        row.cells[1],
                        row.cells[1].text,
                        "Live call intelligence, AI dispatch support, supervisor monitoring, QA — assistive alongside CAD, not a CAD replacement",
                    )
            if texts and "commsupervisor" in texts[0]:
                replace_cell_text(row.cells[0], "supervisor / commsupervisor", "supervisor")
            if texts and texts[0] == "campus_security":
                replace_cell_text(
                    row.cells[2] if len(row.cells) > 2 else row.cells[0],
                    row.cells[2].text if len(row.cells) > 2 else "",
                    "Campus-scoped incident queue, patrol coordination, two-way reporter messaging",
                )
            if any("FedRAMP" in t and "SOC 2" in t for t in texts):
                for cell in row.cells:
                    if "FedRAMP" in cell.text:
                        replace_cell_text(
                            cell,
                            cell.text,
                            "AWS inherited controls (SOC 2 / ISO / FedRAMP on AWS services). Rapid Cortex application: no SOC 2 Type II report claimed; HECVAT and alignment documentation will be provided.",
                        )
            if texts and texts[0] == "AI intelligence processing (RapidIQ): $40,000/year":
                replace_cell_text(row.cells[0], texts[0], "AI incident intelligence (AWS Bedrock): $40,000/year")

    # Expand UM-019 table with campus counselor/faculty if the table is the role table
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[:3] == ["RC Role", "IU Equivalent", "Key Permissions"]:
            existing = {row.cells[0].text.strip() for row in table.rows}
            extras = [
                ("campus_admin", "Campus Administrator", "Users, QR/NFC locations, campus policy, Clery-adjacent configuration"),
                ("campus_supervisor", "Campus Supervisor", "Ops oversight, escalations, cross-building campus view"),
                ("campus_counselor", "Counselor / CARE team", "Welfare-sensitive incidents; not a uniformed security queue"),
                ("campus_faculty", "Faculty / limited reporter", "Limited campus reporting and status; no SOC administration"),
                ("venue_admin / venue_supervisor", "Athletic / venue command", "Memorial Stadium / Assembly Hall game-day console"),
            ]
            # Fix campus_security row label if needed
            for row in table.rows:
                if row.cells[0].text.strip() == "agencyadmin":
                    row.cells[1].paragraphs[0]  # keep
                    set_run_text(row.cells[1].paragraphs[0], "Agency / campus admin (PSAP or campus tenant)")
            for role, equiv, perms in extras:
                if role in existing:
                    continue
                row = table.add_row()
                set_run_text(row.cells[0].paragraphs[0], role)
                set_run_text(row.cells[1].paragraphs[0], equiv)
                set_run_text(row.cells[2].paragraphs[0], perms)

    # Second pass for any leftover phrases
    leftovers = {
        "40-plus languages": "100-plus written language codes",
        "40-plus": "100-plus",
        "RapidIQ AI": "Rapid Cortex incident intelligence",
        "RapidIQ": "Rapid Cortex incident intelligence",
        "commits to compliance": "will align controls under an IU DPA",
        "CJIS certified": "CJIS-aligned",
        "CJIS-compliant": "CJIS-aligned",
    }
    # Avoid over-replacing RapidIQ if already fixed; still catch stragglers
    for p in doc.paragraphs:
        original = p.text
        new = original
        for old, repl in leftovers.items():
            if old in new:
                new = new.replace(old, repl)
        if new != original:
            set_run_text(p, new)
    for table in walk_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    original = p.text
                    new = original
                    for old, repl in leftovers.items():
                        if old in new:
                            new = new.replace(old, repl)
                    if new != original:
                        set_run_text(p, new)

    doc.save(str(OUT_DOCX))


def main() -> None:
    if not SRC_XLSX.exists() or not SRC_DOCX.exists():
        raise SystemExit("Source IU files not found in Downloads")
    align_excel()
    align_docx()
    print("Wrote", OUT_XLSX)
    print("Wrote", OUT_DOCX)


if __name__ == "__main__":
    main()
