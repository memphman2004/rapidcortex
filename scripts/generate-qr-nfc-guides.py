#!/usr/bin/env python3
"""Generate Rapid Cortex QR/NFC field guides (no third-party NFC apps).

Outputs:
  Rapid Cortex Internal Docs/RC_NFC_QR_Setup_Guide.pdf
  Rapid Cortex Internal Docs/Product Usage/RC_NFC_QR_Setup_Guide.pdf
  Rapid Cortex Internal Docs/RC_NFC_Tag_Installation_Guide.pdf
  Rapid Cortex Internal Docs/Product Usage/RC_NFC_Tag_Installation_Guide.pdf

Also patches RC_NFC_Tag_Installation_Guide.docx to remove NFC Tools.
"""

from __future__ import annotations

from pathlib import Path

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "Rapid Cortex Internal Docs"
USAGE = DOCS / "Product Usage"

# Brand
NAVY = colors.HexColor("#0B1220")
NAVY2 = colors.HexColor("#111827")
HEADER = colors.HexColor("#0D1B2A")
BLUE = colors.HexColor("#2563EB")
BLUE_DK = colors.HexColor("#1D4ED8")
SKY = colors.HexColor("#38BDF8")
GREEN = colors.HexColor("#16A34A")
GREEN_DK = colors.HexColor("#166534")
AMBER = colors.HexColor("#D97706")
RED = colors.HexColor("#DC2626")
SLATE = colors.HexColor("#94A3B8")
WHITE = colors.HexColor("#F8FAFC")
CARD = colors.HexColor("#152238")
TIP = colors.HexColor("#0F2744")
BONUS = colors.HexColor("#0F2A1C")
TROUBLE = colors.HexColor("#2A1216")
PAGE_W, PAGE_H = letter
MARGIN = 0.42 * inch


def rounded_rect(c, x, y, w, h, r, fill, stroke=None, sw=0.6):
    c.saveState()
    c.setFillColor(fill)
    if stroke:
        c.setStrokeColor(stroke)
        c.setLineWidth(sw)
        c.roundRect(x, y, w, h, r, fill=1, stroke=1)
    else:
        c.roundRect(x, y, w, h, r, fill=1, stroke=0)
    c.restoreState()


def wrap_text(c, text, font, size, max_w):
    words = text.split()
    lines, cur = [], ""
    for w in words:
        trial = f"{cur} {w}".strip()
        if c.stringWidth(trial, font, size) <= max_w:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = w
    if cur:
        lines.append(cur)
    return lines or [text]


def draw_header(c, badge):
    rounded_rect(c, 0, PAGE_H - 52, PAGE_W, 52, 0, HEADER)
    c.setFillColor(BLUE)
    c.rect(0, PAGE_H - 54, PAGE_W, 2.2, fill=1, stroke=0)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 13)
    c.drawString(MARGIN, PAGE_H - 24, "Rapid Cortex")
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 7.5)
    c.drawString(MARGIN, PAGE_H - 38, "Intelligence at the Speed of Response")
    bw = c.stringWidth(badge, "Helvetica-Bold", 8) + 16
    rounded_rect(c, PAGE_W - MARGIN - bw, PAGE_H - 38, bw, 18, 4, BLUE)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 8)
    c.drawRightString(PAGE_W - MARGIN - 8, PAGE_H - 33, badge)


def draw_footer(c, page, total=2):
    c.setFillColor(HEADER)
    c.rect(0, 0, PAGE_W, 28, fill=1, stroke=0)
    c.setFillColor(BLUE)
    c.rect(0, 28, PAGE_W, 1.6, fill=1, stroke=0)
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 7)
    c.drawString(MARGIN, 11, "Rapid Cortex  ·  support@rapidcortex.us  ·  app.rapidcortex.us")
    c.drawRightString(PAGE_W - MARGIN, 11, f"Page {page} of {total}")


def numbered_step(c, x, y, n, title, body, accent, max_w):
    rounded_rect(c, x, y - 3, 14, 14, 7, accent)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 8)
    c.drawCentredString(x + 7, y, str(n))
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(x + 20, y, title)
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 7.2)
    lines = wrap_text(c, body, "Helvetica", 7.2, max_w - 22)
    yy = y - 12
    for line in lines:
        c.drawString(x + 20, yy, line)
        yy -= 9.2
    return yy - 6


def tool_badge(c, x, y, w, h, title, sub, fill):
    rounded_rect(c, x, y, w, h, 6, fill)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(x + 10, y + h - 16, title)
    c.setFont("Helvetica", 7)
    c.setFillColor(colors.HexColor("#E2E8F0"))
    lines = wrap_text(c, sub, "Helvetica", 7, w - 18)
    yy = y + h - 28
    for line in lines:
        c.drawString(x + 10, yy, line)
        yy -= 9


def phone_frame(c, x, y, w, h, caption, inner_draw):
    rounded_rect(c, x, y, w, h, 10, colors.HexColor("#020617"), colors.HexColor("#334155"), 1)
    rounded_rect(c, x + 6, y + 16, w - 12, h - 28, 7, colors.HexColor("#0B1220"))
    inner_draw(c, x + 10, y + 20, w - 20, h - 36)
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 6.2)
    c.drawCentredString(x + w / 2, y + 5, caption)


def draw_setup_page1(c):
    draw_header(c, "QR & NFC SETUP GUIDE")
    c.setFillColor(NAVY)
    c.rect(0, 28, PAGE_W, PAGE_H - 80, fill=1, stroke=0)

    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(PAGE_W / 2, PAGE_H - 76, "Creating QR Codes and Programming NFC Tags")
    c.setFillColor(SKY)
    c.setFont("Helvetica", 8)
    c.drawCentredString(
        PAGE_W / 2,
        PAGE_H - 92,
        "Everything in the Rapid Cortex ecosystem — no third-party tools required",
    )

    gap = 8
    bw = (PAGE_W - 2 * MARGIN - 2 * gap) / 3
    by = PAGE_H - 148
    tool_badge(c, MARGIN, by, bw, 42, "Browser", "Create & download QR codes", BLUE)
    tool_badge(
        c,
        MARGIN + bw + gap,
        by,
        bw,
        42,
        "RC Mobile App",
        "Program NFC tags  ·  Scan QR codes",
        colors.HexColor("#1E3A5F"),
    )
    tool_badge(
        c,
        MARGIN + 2 * (bw + gap),
        by,
        bw,
        42,
        "RC Dashboard",
        "View analytics & manage codes",
        GREEN_DK,
    )

    col_w = (PAGE_W - 2 * MARGIN - 10) / 2
    left_x = MARGIN
    right_x = MARGIN + col_w + 10
    col_top = PAGE_H - 162
    col_h = 318
    rounded_rect(c, left_x, col_top - col_h, col_w, col_h, 8, CARD, colors.HexColor("#1E3A8A"), 0.8)
    rounded_rect(c, right_x, col_top - col_h, col_w, col_h, 8, CARD, colors.HexColor("#164E63"), 0.8)

    c.setFillColor(SKY)
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(left_x + 12, col_top - 18, "Creating QR Codes")
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 7)
    c.drawString(left_x + 12, col_top - 30, "Browser — app.rapidcortex.us")

    c.setFillColor(colors.HexColor("#7DD3FC"))
    c.setFont("Helvetica-Bold", 9.5)
    c.drawString(right_x + 12, col_top - 18, "Programming NFC Tags")
    c.setFillColor(SLATE)
    c.setFont("Helvetica", 7)
    c.drawString(right_x + 12, col_top - 30, "RC Mobile App — iOS & Android")

    qr_steps = [
        ("Log into the RC web app", "Go to app.rapidcortex.us in any browser."),
        ("Go to QR & NFC", "Sidebar → select your agency → QR & NFC."),
        ("Tap “+ New QR / NFC Code”", "Enter location name, zone, report type, vertical."),
        ("Click “Download PNG”", "QR code image downloads to your computer."),
        ("Print and attach the QR code", "Minimum 2″ on wall signs."),
    ]
    y = col_top - 48
    for i, (t, b) in enumerate(qr_steps, 1):
        y = numbered_step(c, left_x + 10, y, i, t, b, BLUE, col_w - 28)

    nfc_steps = [
        ("Download the RC Mobile App", "App Store or Google Play — search “Rapid Cortex”."),
        ("Sign in with your RC account", "Same credentials as the web app."),
        ("Open QR & NFC", "Select your agency, then open the location list."),
        ("Tap the location to program", "e.g. “Main Lobby” or “Library 3rd Floor”."),
        ("Tap “Program NFC Tag”", "App activates NFC and waits for a tag."),
        ("Hold the NTAG213 to the phone", "Keep still — write takes about 2 seconds."),
        ("Tag programmed — test it", "Tap the tag to confirm the report form opens."),
    ]
    y = col_top - 48
    for i, (t, b) in enumerate(nfc_steps, 1):
        y = numbered_step(c, right_x + 10, y, i, t, b, colors.HexColor("#0EA5E9"), col_w - 28)

    tip_h = 36
    rounded_rect(c, left_x + 8, col_top - col_h + 8, col_w - 16, tip_h, 5, TIP)
    c.setFillColor(SKY)
    c.setFont("Helvetica-Bold", 7)
    c.drawString(left_x + 16, col_top - col_h + 30, "Tip")
    c.setFillColor(colors.HexColor("#CBD5E1"))
    c.setFont("Helvetica", 6.6)
    c.drawString(left_x + 16, col_top - col_h + 18, "QR codes work on every smartphone — no app needed.")
    c.drawString(left_x + 16, col_top - col_h + 9, "Min 1.5″ desk  ·  2″ wall  ·  3″ hallways")

    rounded_rect(c, right_x + 8, col_top - col_h + 8, col_w - 16, tip_h, 5, BONUS)
    c.setFillColor(colors.HexColor("#4ADE80"))
    c.setFont("Helvetica-Bold", 7)
    c.drawString(right_x + 16, col_top - col_h + 30, "Bonus")
    c.setFillColor(colors.HexColor("#CBD5E1"))
    c.setFont("Helvetica", 6.6)
    c.drawString(right_x + 16, col_top - col_h + 18, "The RC app can also scan and verify QR codes.")
    c.drawString(right_x + 16, col_top - col_h + 9, "Confirm a printed code works before mounting.")

    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(MARGIN, 196, "App screens — NFC tag programming flow")

    phones_y = 40
    phone_h = 148
    phone_w = (PAGE_W - 2 * MARGIN - 24) / 4
    labels = [
        "1  Location → Program NFC Tag",
        "2  Hold tag to phone",
        "3  Success — tag is live",
        "4  Scan QR to verify",
    ]

    def s1(c, x, y, w, h):
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 6.5)
        c.drawString(x + 4, y + h - 12, "Main Lobby")
        c.setFillColor(SLATE)
        c.setFont("Helvetica", 5.4)
        c.drawString(x + 4, y + h - 22, "app.rapidcortex.us/report/…")
        for i, (label, fill) in enumerate(
            [
                ("Scan QR Code", BLUE),
                ("Program NFC Tag", colors.HexColor("#0EA5E9")),
                ("View Analytics", colors.HexColor("#334155")),
            ]
        ):
            rounded_rect(c, x + 4, y + h - 48 - i * 22, w - 8, 18, 3, fill)
            c.setFillColor(WHITE)
            c.setFont("Helvetica-Bold", 5.8)
            c.drawCentredString(x + w / 2, y + h - 42 - i * 22, label)

    def s2(c, x, y, w, h):
        c.setFillColor(SKY)
        c.setFont("Helvetica-Bold", 6.5)
        c.drawCentredString(x + w / 2, y + h - 14, "Program NFC Tag")
        c.setStrokeColor(BLUE)
        c.setLineWidth(1.4)
        c.circle(x + w / 2, y + h / 2 + 4, 16, fill=0, stroke=1)
        c.setStrokeColor(colors.HexColor("#38BDF8"))
        c.setLineWidth(0.7)
        c.circle(x + w / 2, y + h / 2 + 4, 22, fill=0, stroke=1)
        c.setFillColor(SLATE)
        c.setFont("Helvetica", 5.4)
        c.drawCentredString(x + w / 2, y + 28, "iPhone: near top edge")
        c.drawCentredString(x + w / 2, y + 18, "Android: center back")
        c.drawCentredString(x + w / 2, y + 8, "Keep tag still…")

    def s3(c, x, y, w, h):
        rounded_rect(c, x + 4, y + h - 28, w - 8, 20, 4, GREEN)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 6.2)
        c.drawCentredString(x + w / 2, y + h - 21, "Tag programmed")
        c.setFillColor(colors.HexColor("#BBF7D0"))
        c.setFont("Helvetica", 5.6)
        c.drawCentredString(x + w / 2, y + h - 44, "Main Lobby tag is live.")
        c.drawCentredString(x + w / 2, y + h - 54, "Tap the tag to test it.")
        rounded_rect(c, x + 8, y + 12, w - 16, 16, 3, BLUE)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 5.5)
        c.drawCentredString(x + w / 2, y + 17, "Write Another")

    def s4(c, x, y, w, h):
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 6.5)
        c.drawCentredString(x + w / 2, y + h - 14, "Scan QR Code")
        c.setStrokeColor(SKY)
        c.setLineWidth(1)
        c.rect(x + 18, y + 36, w - 36, w - 36, fill=0, stroke=1)
        rounded_rect(c, x + 6, y + 10, w - 12, 18, 3, GREEN)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 5.6)
        c.drawCentredString(x + w / 2, y + 16, "QR code valid")

    drawers = [s1, s2, s3, s4]
    for i, draw in enumerate(drawers):
        px = MARGIN + i * (phone_w + 8)
        phone_frame(c, px, phones_y, phone_w, phone_h, labels[i], draw)

    draw_footer(c, 1)


def draw_setup_page2(c):
    draw_header(c, "QR & NFC SETUP GUIDE")
    c.setFillColor(NAVY)
    c.rect(0, 28, PAGE_W, PAGE_H - 80, fill=1, stroke=0)

    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 15)
    c.drawCentredString(PAGE_W / 2, PAGE_H - 76, "What You Need  ·  Tips  ·  Troubleshooting")

    rounded_rect(c, MARGIN, PAGE_H - 250, PAGE_W - 2 * MARGIN, 158, 8, CARD)
    c.setFillColor(SKY)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + 14, PAGE_H - 110, "What you need")

    needs = [
        ("RC admin account (web)", "app.rapidcortex.us · Agency Admin or higher · same account as the mobile app"),
        ("RC Mobile App", "App Store or Google Play · sign in with your RC credentials"),
        ("NTAG213 NFC sticker tags", "Search “NTAG213 NFC stickers” on Amazon · ~$15–20 per 100 · round or square · white"),
        ("Printer for QR codes", "Standard office printer · laser preferred · label paper or cardstock"),
    ]
    y = PAGE_H - 132
    for title, body in needs:
        rounded_rect(c, MARGIN + 12, y - 8, 8, 8, 2, BLUE)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 8)
        c.drawString(MARGIN + 28, y - 7, title)
        c.setFillColor(SLATE)
        c.setFont("Helvetica", 7)
        c.drawString(MARGIN + 28, y - 19, body)
        y -= 32

    col_w = (PAGE_W - 2 * MARGIN - 10) / 2
    tips_top = PAGE_H - 266
    tips_h = 268
    rounded_rect(c, MARGIN, tips_top - tips_h, col_w, tips_h, 8, CARD)
    rounded_rect(c, MARGIN + col_w + 10, tips_top - tips_h, col_w, tips_h, 8, CARD)

    c.setFillColor(SKY)
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + 12, tips_top - 18, "Tips for best results")
    tips = [
        ("Create one code per physical location", "Never reuse URLs — you need per-location data to know where reports come from."),
        ("Label NFC tags before sticking", "Write the location name on the tag back with a marker before peeling."),
        ("Test before mounting", "Always tap the NFC tag and scan the QR code before attaching to a sign."),
        ("QR codes work on all phones", "NFC needs a newer phone. Put both QR and NFC on the same sign for full coverage."),
        ("RC app scans QR codes too", "Use the RC mobile app to verify any printed QR code before you mount it."),
    ]
    y = tips_top - 36
    for title, body in tips:
        c.setFillColor(BLUE)
        c.circle(MARGIN + 18, y + 3, 3.2, fill=1, stroke=0)
        c.setFillColor(WHITE)
        c.setFont("Helvetica-Bold", 7.6)
        c.drawString(MARGIN + 28, y, title)
        c.setFillColor(SLATE)
        c.setFont("Helvetica", 6.8)
        for line in wrap_text(c, body, "Helvetica", 6.8, col_w - 44):
            y -= 10
            c.drawString(MARGIN + 28, y, line)
        y -= 16

    c.setFillColor(colors.HexColor("#FCA5A5"))
    c.setFont("Helvetica-Bold", 10)
    c.drawString(MARGIN + col_w + 22, tips_top - 18, "Troubleshooting")
    issues = [
        ("NFC tag not detected", "Turn NFC on (Settings → Connected devices → NFC). Move the tag slowly around the phone back to find the antenna."),
        ("Wrong page opens", "The tag has an old URL. Open the location in the RC app and tap Program NFC Tag again. NTAG213 tags can be overwritten."),
        ("QR code won’t scan", "Check print size (min 1.5″). Clean smudges. Do not stretch or distort the PNG."),
        ("NFC works but QR doesn’t", "Re-download the QR PNG from the RC dashboard. Print larger or at higher quality."),
        ("App says “Write failed”", "Enable NFC in phone settings. Hold the tag still — movement during write causes failures."),
    ]
    y = tips_top - 40
    ix = MARGIN + col_w + 18
    for title, body in issues:
        box_h = 42
        rounded_rect(c, ix, y - box_h + 12, col_w - 16, box_h, 4, TROUBLE)
        c.setFillColor(colors.HexColor("#FECACA"))
        c.setFont("Helvetica-Bold", 7.2)
        c.drawString(ix + 8, y, title)
        c.setFillColor(colors.HexColor("#FECACA"))
        c.setFont("Helvetica", 6.4)
        ly = y - 11
        for line in wrap_text(c, body, "Helvetica", 6.4, col_w - 34):
            c.drawString(ix + 8, ly, line)
            ly -= 8.5
        y -= 46

    qy = 40
    qh = 118
    rounded_rect(c, MARGIN, qy, PAGE_W - 2 * MARGIN, qh, 8, CARD)
    c.setFillColor(WHITE)
    c.setFont("Helvetica-Bold", 9)
    c.drawString(MARGIN + 12, qy + qh - 16, "Quick reference")
    cols = [
        (
            "QR CODES (Browser)",
            [
                "app.rapidcortex.us → QR & NFC",
                "+ New QR / NFC Code → fill details",
                "Download PNG",
                "Print and mount",
            ],
        ),
        (
            "NFC TAGS (RC Mobile App)",
            [
                "Open RC app → QR & NFC",
                "Tap location → Program NFC Tag",
                "Hold NTAG213 tag to phone",
                "Test · then mount the sign",
            ],
        ),
        (
            "VERIFY (RC Mobile App)",
            [
                "Open the RC app",
                "Tap Scan QR Code",
                "Point at the printed QR code",
                "Confirm the report form opens",
            ],
        ),
    ]
    cw = (PAGE_W - 2 * MARGIN - 24) / 3
    for i, (title, steps) in enumerate(cols):
        x = MARGIN + 12 + i * (cw + 6)
        c.setFillColor(SKY)
        c.setFont("Helvetica-Bold", 7)
        c.drawString(x, qy + qh - 32, title)
        c.setFillColor(SLATE)
        c.setFont("Helvetica", 6.6)
        yy = qy + qh - 46
        for n, step in enumerate(steps, 1):
            c.drawString(x, yy, f"{n}.  {step}")
            yy -= 12

    draw_footer(c, 2)


def write_setup_guide(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=letter)
    c.setTitle("Rapid Cortex QR & NFC Setup Guide")
    c.setAuthor("Rapid Cortex")
    draw_setup_page1(c)
    c.showPage()
    draw_setup_page2(c)
    c.save()


# ── 6-page NFC installation guide (light, internal) ──────────────────────────

LIGHT_NAVY = colors.HexColor("#0F172A")
LIGHT_RULE = colors.HexColor("#CBD5E1")
BODY = colors.HexColor("#1E293B")
MUTED = colors.HexColor("#475569")


def nfc_styles():
    return {
        "h1": ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=22, textColor=LIGHT_NAVY, leading=26, spaceAfter=4),
        "sub": ParagraphStyle("sub", fontName="Helvetica", fontSize=10, textColor=MUTED, leading=13, spaceAfter=10),
        "h2": ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=13, textColor=LIGHT_NAVY, leading=16, spaceBefore=8, spaceAfter=4),
        "h3": ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=10.5, textColor=BLUE_DK, leading=13, spaceBefore=4, spaceAfter=3),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=9.5, textColor=BODY, leading=13, spaceAfter=6),
        "small": ParagraphStyle("small", fontName="Helvetica", fontSize=8.5, textColor=MUTED, leading=11),
        "callout": ParagraphStyle("callout", fontName="Helvetica", fontSize=9, textColor=BODY, leading=12),
        "th": ParagraphStyle("th", fontName="Helvetica-Bold", fontSize=8.5, textColor=WHITE, leading=11),
        "td": ParagraphStyle("td", fontName="Helvetica", fontSize=8.5, textColor=BODY, leading=11),
        "tdb": ParagraphStyle("tdb", fontName="Helvetica-Bold", fontSize=8.5, textColor=LIGHT_NAVY, leading=11),
        "center": ParagraphStyle("center", fontName="Helvetica", fontSize=8.5, textColor=MUTED, alignment=TA_CENTER, leading=11),
        "stepn": ParagraphStyle("stepn", fontName="Helvetica-Bold", fontSize=11, textColor=WHITE, alignment=TA_CENTER, leading=13),
        "stept": ParagraphStyle("stept", fontName="Helvetica-Bold", fontSize=9.5, textColor=LIGHT_NAVY, leading=12),
        "stepb": ParagraphStyle("stepb", fontName="Helvetica", fontSize=9, textColor=BODY, leading=12),
    }


def nfc_header_footer(canvas_obj, doc):
    canvas_obj.saveState()
    canvas_obj.setFillColor(LIGHT_NAVY)
    canvas_obj.rect(0, PAGE_H - 36, PAGE_W, 36, fill=1, stroke=0)
    canvas_obj.setFillColor(BLUE)
    canvas_obj.rect(0, PAGE_H - 38, PAGE_W, 2, fill=1, stroke=0)
    canvas_obj.setFillColor(WHITE)
    canvas_obj.setFont("Helvetica-Bold", 8)
    canvas_obj.drawString(0.6 * inch, PAGE_H - 22, "RAPID CORTEX")
    canvas_obj.setFont("Helvetica", 8)
    canvas_obj.setFillColor(SLATE)
    canvas_obj.drawString(1.55 * inch, PAGE_H - 22, "|  NFC Tag Installation Guide  ·  Confidential — Internal Use Only")
    canvas_obj.setFillColor(colors.HexColor("#F1F5F9"))
    canvas_obj.rect(0, 0, PAGE_W, 28, fill=1, stroke=0)
    canvas_obj.setFillColor(MUTED)
    canvas_obj.setFont("Helvetica", 7.5)
    canvas_obj.drawString(0.6 * inch, 11, "www.rapidcortex.us  ·  support@rapidcortex.us")
    canvas_obj.drawRightString(PAGE_W - 0.6 * inch, 11, f"Page {doc.page}")
    canvas_obj.restoreState()


def step_row(n, title, body, S):
    num = Paragraph(str(n), S["stepn"])
    text = [Paragraph(title, S["stept"]), Paragraph(body, S["stepb"])]
    inner = Table([[text]], colWidths=[6.3 * inch])
    inner.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 2)]))
    t = Table([[num, inner]], colWidths=[0.38 * inch, 6.42 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), BLUE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 2),
                ("RIGHTPADDING", (0, 0), (0, 0), 2),
                ("TOPPADDING", (0, 0), (0, 0), 4),
                ("BOTTOMPADDING", (0, 0), (0, 0), 4),
                ("LEFTPADDING", (1, 0), (1, 0), 8),
                ("TOPPADDING", (1, 0), (1, 0), 1),
                ("BOTTOMPADDING", (1, 0), (1, 0), 8),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F8FAFC")),
                ("BOX", (0, 0), (-1, -1), 0.3, LIGHT_RULE),
            ]
        )
    )
    return t


def callout(kind, text, S):
    fill = {
        "NOTE": colors.HexColor("#EFF6FF"),
        "TIP": colors.HexColor("#ECFDF5"),
        "WARNING": colors.HexColor("#FEF3C7"),
        "CAUTION": colors.HexColor("#FEF2F2"),
    }[kind]
    border = {
        "NOTE": BLUE,
        "TIP": GREEN,
        "WARNING": AMBER,
        "CAUTION": RED,
    }[kind]
    label = Paragraph(f"<b>{kind}</b>", ParagraphStyle("k", fontName="Helvetica-Bold", fontSize=8, textColor=border, leading=10))
    body = Paragraph(text, S["callout"])
    t = Table([[label, body]], colWidths=[0.9 * inch, 5.9 * inch])
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("BOX", (0, 0), (-1, -1), 1, border),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return t


def write_nfc_install_guide(path: Path) -> None:
    S = nfc_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.5 * inch,
        title="Rapid Cortex NFC Tag Installation Guide",
        author="Rapid Cortex",
    )
    story = []
    story.append(Paragraph("NFC Tag Installation Guide", S["h1"]))
    story.append(Paragraph("Rapid Cortex Campus · Rapid Cortex Venue", S["sub"]))
    story.append(Paragraph("Step-by-step guide to programming and attaching NFC tags using only Rapid Cortex — no third-party NFC apps.", S["body"]))

    pills = Table(
        [[Paragraph(x, S["center"]) for x in ["15 minutes", "No technical skills needed", "Works on any smartphone"]]],
        colWidths=[2.3 * inch, 2.3 * inch, 2.3 * inch],
    )
    pills.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF6FF")),
                ("BOX", (0, 0), (-1, -1), 0.4, BLUE),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, LIGHT_RULE),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(pills)
    story.append(Spacer(1, 10))
    story.append(Paragraph("Overview", S["h2"]))
    story.append(
        Paragraph(
            "An NFC (Near Field Communication) tag is a small sticker — about the size of a coin — that contains a tiny chip. "
            "When a smartphone is held near it, the phone opens a web page. No camera, no app download, and no typing for the person reporting.",
            S["body"],
        )
    )
    story.append(
        Paragraph(
            "For Rapid Cortex, each NFC tag is programmed with a unique reporting URL tied to your agency and location. "
            "You create the code in the browser, then write the tag with the <b>RC Mobile App</b>. When someone taps the sign, they go directly to the safety reporting form for that exact location.",
            S["body"],
        )
    )
    story.append(callout("NOTE", "NFC tags and QR codes on the same sign point to the same location. They work identically — NFC is faster for newer phones; QR covers every camera.", S))
    story.append(Paragraph("What You Need", S["h2"]))
    story.append(Paragraph("The Rapid Cortex browser and mobile app do the programming. You only buy blank NTAG213 stickers.", S["body"]))

    header = [Paragraph(x, S["th"]) for x in ["Item", "Details", "Where to get it"]]
    rows = [
        ["RC admin account", "Agency Admin or higher · same login as the mobile app", "app.rapidcortex.us"],
        ["RC Mobile App", "Writes NFC tags and can scan/verify QR codes · no NFC Tools or other apps", "App Store / Google Play — search Rapid Cortex"],
        ["NTAG213 NFC stickers", "Pack of 100 · round or square · white finish", "Amazon — search “NTAG213 NFC stickers” (~$15–20 / 100)"],
        ["Smartphone with NFC", "iPhone 7 or newer · most Android phones since 2014", "You already have this"],
        ["Permanent marker (optional)", "Label tag backs with the location name", "Any office supply store"],
        ["Isopropyl alcohol wipes (optional)", "Clean the sign surface before sticking", "Amazon or pharmacy"],
    ]
    data = [header] + [[Paragraph(a, S["tdb"]), Paragraph(b, S["td"]), Paragraph(c, S["td"])] for a, b, c in rows]
    tbl = Table(data, colWidths=[1.8 * inch, 2.8 * inch, 2.3 * inch])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_NAVY),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("GRID", (0, 0), (-1, -1), 0.3, LIGHT_RULE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(tbl)
    story.append(Spacer(1, 8))
    story.append(callout("TIP", "NTAG213 is the most compatible chip type. NTAG215 / NTAG216 work but cost more for no benefit in this use case.", S))

    story.append(Paragraph("1  Create the location in the browser", S["h2"]))
    story.append(Paragraph("Each physical sign location gets its own unique code. You generate it in the Rapid Cortex web app — you do not paste URLs into any other app.", S["body"]))
    for n, title, body in [
        (1, "Log in to the Rapid Cortex web app", "Go to app.rapidcortex.us and sign in with your admin account."),
        (2, "Open QR & NFC", "From the sidebar, select your agency, then click QR & NFC. RC platform admins go to RC Admin → Agencies → [Agency] → QR & NFC."),
        (3, "Tap “+ New QR / NFC Code”", "Enter the name and location of this sign (e.g., “McKinley Hall — 3rd Floor”), select the report type and vertical, and save."),
        (4, "Leave the URL in Rapid Cortex", "The RC Mobile App writes the correct URL for you. You do not need to copy it into NFC Tools or any other writer."),
    ]:
        story.append(step_row(n, title, body, S))
    story.append(Spacer(1, 6))
    story.append(callout("WARNING", "Create a separate code for each physical sign location. Do not reuse the same URL across multiple locations — you will not be able to tell which location a report came from.", S))

    story.append(Paragraph("2  Program the NFC tag in the RC Mobile App", S["h2"]))
    story.append(Paragraph("This takes about 10 seconds per tag once you know the steps. You only need to do this once per tag. The RC app writes the URL — no third-party NFC utility.", S["body"]))
    for n, title, body in [
        (1, "Download the RC Mobile App", "Search “Rapid Cortex” in the App Store (iOS) or Google Play (Android). Sign in with the same account you use on the web."),
        (2, "Open QR & NFC", "Select your agency, then open the location you just created."),
        (3, "Tap Program NFC Tag", "The app activates NFC and waits for a tag. You do not paste a URL."),
        (4, "Hold an NTAG213 to the back of the phone", "iPhone: near the top edge. Android: center back. Hold still for about 2 seconds until the app shows Tag programmed successfully."),
        (5, "Write another tag if needed", "Use Write Another in the app for extra tags at the same location. Each tap is counted separately."),
    ]:
        story.append(step_row(n, title, body, S))
    story.append(Spacer(1, 6))
    story.append(callout("TIP", "If the write fails, move the tag slowly around the back of the phone until you find the NFC antenna. On iPhone it is near the top edge. On Android it is usually in the center.", S))

    story.append(Paragraph("3  Test the tag before attaching", S["h2"]))
    story.append(Paragraph("Always test before mounting. Testing takes 30 seconds and prevents pulling a tag out from under a sign.", S["body"]))
    for n, title, body in [
        (1, "Tap the programmed tag with a locked-awake phone", "Screen on, no app required for the person reporting. Hold the tag to the NFC antenna."),
        (2, "Confirm the Rapid Cortex report form opens", "You should see the safety reporting form for this location and vertical."),
        (3, "Optional: verify the printed QR in the RC app", "Use Scan QR Code in the RC Mobile App to confirm the printed PNG before you mount."),
        (4, "Check the dashboard", "In the RC web app, confirm the NFC tap count for this code increased by 1."),
    ]:
        story.append(step_row(n, title, body, S))
    story.append(Spacer(1, 6))
    story.append(callout("CAUTION", "If the form does not open: confirm NFC is enabled (Settings → NFC), rewrite the tag from the RC Mobile App (Program NFC Tag), and hold the tag at the antenna location.", S))

    story.append(Paragraph("4  Attach the tag to the sign", S["h2"]))
    for n, title, body in [
        (1, "Label the tag (recommended)", "Before peeling, write the location name on the tag (e.g., “MH-3F”). Helps if a tag ever comes loose."),
        (2, "Clean the mounting surface (optional)", "Wipe the back of the sign with isopropyl alcohol and let dry 30 seconds."),
        (3, "Peel and press", "Center the tag on the back of the sign. Press firmly 10–15 seconds."),
        (4, "Mount the sign", "Install in the final location. Anyone with a smartphone can tap the front of the sign to report."),
        (5, "Final tap test", "Hold your phone to the front of the mounted sign. Confirm the form still opens."),
    ]:
        story.append(step_row(n, title, body, S))

    story.append(Paragraph("Sign material compatibility", S["h2"]))
    story.append(Paragraph("NFC works differently depending on what the sign is made of.", S["body"]))
    mat_header = [Paragraph(x, S["th"]) for x in ["Material", "NFC works?", "Notes"]]
    mats = [
        ["Acrylic / Plexiglass", "Yes", "Best option. Tag reads clearly through any thickness."],
        ["PVC / Foam board", "Yes", "Reads through up to 1/2 inch thickness."],
        ["Plastic", "Yes", "Works on any non-metallic plastic."],
        ["Wood", "Yes", "Works through standard sign thickness."],
        ["Glass", "Yes", "Reads through glass. Stick tag to inner surface."],
        ["Aluminum (face)", "Partial", "Aluminum blocks NFC. Mount tag on the face (not back) in a corner, or use a mounting block."],
        ["Solid metal back", "No", "Metal blocks the signal. Use a non-metallic plate between tag and metal, or attach nearby."],
    ]
    md = [mat_header] + [[Paragraph(a, S["tdb"]), Paragraph(b, S["td"]), Paragraph(c, S["td"])] for a, b, c in mats]
    mt = Table(md, colWidths=[1.8 * inch, 1.1 * inch, 4.0 * inch])
    mt.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_NAVY),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("GRID", (0, 0), (-1, -1), 0.3, LIGHT_RULE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(mt)

    story.append(Paragraph("5  Troubleshooting", S["h2"]))
    t_header = [Paragraph(x, S["th"]) for x in ["Problem", "Solution"]]
    troubles = [
        ["Phone does not detect the tag", "Enable NFC: iPhone Control Center / Android Settings → Connected devices → NFC. Move the tag slowly across the phone back."],
        ["Wrong page opens", "Rewrite the tag from the RC Mobile App (Program NFC Tag). NTAG213 tags can be overwritten. Do not use NFC Tools or another writer."],
        ["Tag detected but no page opens", "Rewrite from the RC app so the tag gets a full https:// Rapid Cortex URL."],
        ["Tag stopped working after mounting", "Too close to metal. Move to a non-metallic area or use an on-metal / anti-metal NFC tag."],
        ["Tap count not updating", "Count increments when the reporting form loads. Check connectivity at the sign location."],
        ["Correct form, wrong location name", "Update the zone name in the Rapid Cortex dashboard. No need to rewrite the tag."],
        ["Adhesive came loose", "Clean with isopropyl alcohol, dry fully, and re-stick. For permanent installs, a small amount of clear super glue around the edge."],
        ["App says write failed", "NFC must be on. Hold the tag still. Use an NTAG213. Stay in the RC Mobile App — do not switch to a third-party NFC utility."],
    ]
    td = [t_header] + [[Paragraph(a, S["tdb"]), Paragraph(b, S["td"])] for a, b in troubles]
    tt = Table(td, colWidths=[2.1 * inch, 4.8 * inch])
    tt.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), LIGHT_NAVY),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8FAFC")]),
                ("GRID", (0, 0), (-1, -1), 0.3, LIGHT_RULE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(tt)

    story.append(Paragraph("6  Quick reference card", S["h2"]))
    story.append(Paragraph("Print and keep with your NFC tag supplies.", S["body"]))
    left = Paragraph(
        "<b>Supplies</b><br/>☐ NTAG213 NFC sticker tags<br/>☐ Smartphone with NFC<br/>☐ RC Mobile App (not NFC Tools)<br/>☐ RC admin login<br/>☐ Permanent marker (optional)<br/>☐ Isopropyl wipes (optional)",
        S["td"],
    )
    right = Paragraph(
        "<b>Installation steps</b><br/>1. Create the location at app.rapidcortex.us → QR &amp; NFC<br/>2. Open RC app → that location → Program NFC Tag<br/>3. Hold NTAG213 to the phone until success<br/>4. Test tap — confirm the form opens<br/>5. Optional: Scan QR Code in the RC app<br/>6. Label, clean, peel, stick, mount<br/>7. Final tap test after mounting",
        S["td"],
    )
    q = Table([[left, right]], colWidths=[3.45 * inch, 3.45 * inch])
    q.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8FAFC")),
                ("BOX", (0, 0), (-1, -1), 0.6, BLUE),
                ("LINEAFTER", (0, 0), (0, 0), 0.4, LIGHT_RULE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.append(q)
    story.append(Spacer(1, 16))
    story.append(Paragraph("Need help? Contact Rapid Cortex Support — support@rapidcortex.us · www.rapidcortex.us", S["body"]))

    doc.build(story, onFirstPage=nfc_header_footer, onLaterPages=nfc_header_footer)


def set_cell_text(cell, text: str) -> None:
    paras = list(cell.paragraphs)
    if not paras:
        return
    first = paras[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for para in paras[1:]:
        for run in para.runs:
            run.text = ""


def patch_nfc_docx(path: Path) -> None:
    from docx import Document

    doc = Document(str(path))
    t3 = doc.tables[3]
    set_cell_text(t3.rows[3].cells[0], "RC Mobile App")
    set_cell_text(
        t3.rows[3].cells[1],
        "Writes NFC tags and scans/verifies QR codes · iOS and Android · no NFC Tools or other apps",
    )
    set_cell_text(t3.rows[3].cells[2], "App Store / Google Play — search Rapid Cortex")
    set_cell_text(t3.rows[4].cells[0], "Location created in Rapid Cortex")
    set_cell_text(
        t3.rows[4].cells[1],
        "Create the QR/NFC code in the web app. The RC Mobile App writes the URL — do not paste it into a third-party writer.",
    )
    set_cell_text(t3.rows[4].cells[2], "app.rapidcortex.us → QR & NFC")

    set_cell_text(
        doc.tables[6].rows[4].cells[1],
        "Click “+ New QR / NFC Code”\nEnter the name and location of this sign (e.g., “McKinley Hall — 3rd Floor”), select the report type, and click Save.",
    )
    set_cell_text(
        doc.tables[6].rows[6].cells[1],
        "Do not copy the URL into another app\nThe RC Mobile App writes the correct URL when you tap Program NFC Tag. You do not paste anything into NFC Tools or any other writer.",
    )

    set_cell_text(
        doc.tables[8].rows[0].cells[1],
        "Program the NFC Tag\nUse the Rapid Cortex Mobile App to write the URL to the tag — no third-party tools.",
    )

    t9 = doc.tables[9]
    set_cell_text(
        t9.rows[0].cells[1],
        "Download the RC Mobile App\nSearch “Rapid Cortex” in the App Store (iOS) or Google Play (Android). Sign in with the same account you use on the web.",
    )
    set_cell_text(
        t9.rows[2].cells[1],
        "Open QR & NFC\nSelect your agency, then open the location you created in the web app.",
    )
    set_cell_text(
        t9.rows[4].cells[1],
        "Tap Program NFC Tag\nThe app activates NFC and waits for a tag. You do not paste a URL.",
    )
    set_cell_text(
        t9.rows[6].cells[1],
        "Hold an NTAG213 to the back of your phone\niPhone: near the top edge. Android: center back. Hold still for about 2 seconds.",
    )
    set_cell_text(
        t9.rows[8].cells[1],
        "Wait for Tag programmed successfully\nThen tap the tag to confirm the Rapid Cortex report form opens. Use Write Another for extra tags at the same location.",
    )
    set_cell_text(
        t9.rows[10].cells[1],
        "Stay in the RC app\nDo not switch to NFC Tools or any other NFC utility. Rapid Cortex writes and verifies the tag.",
    )

    set_cell_text(
        doc.tables[14].rows[0].cells[1],
        "If the form does not open: confirm that NFC is enabled on your phone (Settings → NFC), rewrite the tag from the RC Mobile App (Program NFC Tag), and make sure you are holding the tag in the right position.",
    )

    t19 = doc.tables[19]
    set_cell_text(
        t19.rows[2].cells[1],
        "The tag was programmed with the wrong URL. Open the location in the RC Mobile App and tap Program NFC Tag again. You can overwrite an NTAG213 at any time. Do not use NFC Tools.",
    )
    set_cell_text(
        t19.rows[3].cells[1],
        "Rewrite the tag from the RC Mobile App so it receives a full https:// Rapid Cortex URL. Do not use a third-party NFC writer.",
    )

    set_cell_text(
        doc.tables[21].rows[0].cells[0],
        "Supplies Checklist\n☐  NTAG213 NFC sticker tags\n☐  Smartphone with NFC\n☐  RC Mobile App (not NFC Tools)\n☐  RC admin login at app.rapidcortex.us\n☐  Permanent marker (optional)\n☐  Isopropyl wipes (optional)",
    )
    set_cell_text(
        doc.tables[21].rows[0].cells[1],
        "Installation Steps\n1.  Create location: app.rapidcortex.us → QR & NFC\n2.  Open RC app → location → Program NFC Tag\n3.  Hold NTAG213 to the phone until success\n4.  Test tap — confirm form opens\n5.  Optional: Scan QR Code in the RC app\n6.  Label tag with location name\n7.  Clean sign surface\n8.  Peel and stick to back of sign\n9.  Mount sign\n10.  Final tap test after mounting",
    )

    for para in doc.paragraphs:
        if "Amazon" in para.text and "Purchase these items" in para.text:
            for run in para.runs:
                run.text = run.text.replace(
                    "Purchase these items before installation. Everything is available on Amazon.",
                    "Create the location in the Rapid Cortex web app, then program tags with the RC Mobile App. Buy blank NTAG213 stickers (Amazon is fine for tags only).",
                )

    doc.save(str(path))


def copy_to(src: Path, dest: Path) -> None:
    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_bytes(src.read_bytes())


def main() -> None:
    setup = DOCS / "RC_NFC_QR_Setup_Guide.pdf"
    nfc_pdf = DOCS / "RC_NFC_Tag_Installation_Guide.pdf"
    nfc_docx = DOCS / "RC_NFC_Tag_Installation_Guide.docx"

    write_setup_guide(setup)
    write_nfc_install_guide(nfc_pdf)
    if nfc_docx.exists():
        patch_nfc_docx(nfc_docx)

    copy_to(setup, USAGE / "RC_NFC_QR_Setup_Guide.pdf")
    copy_to(nfc_pdf, USAGE / "RC_NFC_Tag_Installation_Guide.pdf")

    downloads = Path.home() / "Downloads" / "RC_NFC_QR_Setup_Guide.pdf"
    try:
        copy_to(setup, downloads)
    except OSError:
        pass

    print(f"Wrote {setup}")
    print(f"Wrote {nfc_pdf}")
    if nfc_docx.exists():
        print(f"Patched {nfc_docx}")


if __name__ == "__main__":
    main()
